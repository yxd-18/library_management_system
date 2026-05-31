from __future__ import annotations

import csv
import io
import os
import shutil
import socket
import subprocess
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = Path(
    r"D:\WeChat_files\WeChat Files\wxid_q8ndugwzao5a22\FileStorage\File\2026-05\图书管理系统研发文档(1)_完善版_修订.docx"
)
OUTPUT_DOC = Path(
    r"D:\WeChat_files\WeChat Files\wxid_q8ndugwzao5a22\FileStorage\File\2026-05\图书管理系统研发文档(1)_完善版_修订_已插入测试图片.docx"
)
FIG_DIR = PROJECT_ROOT / "generated_test_figures"
QT_IMG_DIR = PROJECT_ROOT / "tmp_softeng_images"
DB_CONFIG_PATH = PROJECT_ROOT / "db_config.txt"
PSQL_PATH = Path(r"C:\Program Files\PostgreSQL\10\bin\psql.exe")
FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")

BG = "#f6f2eb"
CARD = "#ffffff"
BORDER = "#d8cec2"
HEADER_BG = "#efe3d3"
HEADER_TEXT = "#6d4325"
TEXT = "#2f261f"
MUTED = "#73695d"
ACCENT = "#20c400"


@dataclass(frozen=True)
class FigureSpec:
    image_path: Path
    caption: str
    note: str
    width_cm: float = 14.6


def load_font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_PATH), size=size)


FONT_TITLE = load_font(28)
FONT_SUBTITLE = load_font(18)
FONT_HEAD = load_font(20)
FONT_BODY = load_font(18)
FONT_PANEL = load_font(22)


def read_db_config() -> dict[str, str]:
    config: dict[str, str] = {
        "host": "localhost",
        "port": "5432",
        "dbname": "library_management_system",
        "user": "postgres",
        "password": "",
    }
    for raw in DB_CONFIG_PATH.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        config[key.strip()] = value.strip()
    return config


def run_copy_query(config: dict[str, str], sql: str) -> list[dict[str, str]]:
    env = os.environ.copy()
    env["PGPASSWORD"] = config["password"]
    env["PGCLIENTENCODING"] = "UTF8"
    command = [
        str(PSQL_PATH),
        "-h",
        config["host"],
        "-p",
        config["port"],
        "-U",
        config["user"],
        "-d",
        config["dbname"],
        "-c",
        f"COPY ({sql}) TO STDOUT WITH CSV HEADER",
    ]
    result = subprocess.run(command, env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=False)
    if result.returncode != 0:
        stderr = result.stderr.decode("utf-8", errors="replace").strip()
        raise RuntimeError(f"SQL execution failed: {stderr or sql}")
    text = result.stdout.decode("utf-8-sig", errors="replace").replace("\r\n", "\n").strip()
    if not text:
        return []
    return list(csv.DictReader(io.StringIO(text)))


def run_scalar(config: dict[str, str], sql: str) -> str:
    rows = run_copy_query(config, sql)
    if not rows:
        return ""
    first = rows[0]
    return str(next(iter(first.values()), ""))


def tcp_reachable(host: str, port: int) -> bool:
    try:
        with socket.create_connection((host, port), timeout=2):
            return True
    except OSError:
        return False


def fit_text(text: str, font: ImageFont.FreeTypeFont, max_width: int) -> str:
    value = "" if text is None else str(text)
    if font.getlength(value) <= max_width:
        return value
    ellipsis = "…"
    candidate = value
    while candidate and font.getlength(candidate + ellipsis) > max_width:
        candidate = candidate[:-1]
    return (candidate + ellipsis) if candidate else ellipsis


def draw_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.FreeTypeFont,
              fill: str, align: str = "left") -> None:
    left, top, right, bottom = box
    display = fit_text(text, font, max_width=max(40, right - left - 24))
    bbox = draw.textbbox((0, 0), display, font=font)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    if align == "center":
        x = left + (right - left - text_w) / 2
    elif align == "right":
        x = right - text_w - 12
    else:
        x = left + 12
    y = top + (bottom - top - text_h) / 2 - 2
    draw.text((x, y), display, font=font, fill=fill)


def render_table_image(
    output: Path,
    title: str,
    subtitle: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    alignments: list[str] | None = None,
) -> None:
    alignments = alignments or ["left"] * len(headers)
    top_pad = 34
    side_pad = 34
    row_h = 54
    table_w = sum(widths)
    title_gap = 24
    subtitle_gap = 16 if subtitle else 0
    header_y = top_pad + 24 + 44 + title_gap + (26 if subtitle else 0) + subtitle_gap
    img_w = table_w + side_pad * 2
    img_h = header_y + row_h * (len(rows) + 1) + 40

    image = Image.new("RGB", (img_w, img_h), BG)
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, img_w, 18), fill=ACCENT)
    draw.rounded_rectangle((18, 28, img_w - 18, img_h - 18), radius=18, fill=CARD, outline=BORDER, width=2)

    draw.text((side_pad, top_pad + 14), title, font=FONT_TITLE, fill=TEXT)
    if subtitle:
        draw.text((side_pad, top_pad + 60), subtitle, font=FONT_SUBTITLE, fill=MUTED)

    x = side_pad
    y = header_y
    current_x = x
    for idx, (header, width) in enumerate(zip(headers, widths)):
        draw.rectangle((current_x, y, current_x + width, y + row_h), fill=HEADER_BG, outline=BORDER, width=2)
        draw_text(draw, (current_x, y, current_x + width, y + row_h), header, FONT_HEAD, HEADER_TEXT, "center")
        current_x += width

    for row_index, row in enumerate(rows, start=1):
        row_y = y + row_h * row_index
        fill = "#fffdfa" if row_index % 2 else "#fbf7f1"
        current_x = x
        for col_index, (cell, width) in enumerate(zip(row, widths)):
            draw.rectangle((current_x, row_y, current_x + width, row_y + row_h), fill=fill, outline=BORDER, width=1)
            draw_text(
                draw,
                (current_x, row_y, current_x + width, row_y + row_h),
                cell,
                FONT_BODY,
                TEXT,
                alignments[col_index],
            )
            current_x += width

    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output)


def add_panel_background(image: Image.Image, label: str, width: int) -> Image.Image:
    padding = 18
    label_h = 48
    inner = image.copy()
    card = Image.new("RGB", (width, inner.height + padding * 2 + label_h), BG)
    draw = ImageDraw.Draw(card)
    draw.rounded_rectangle((8, 8, width - 8, card.height - 8), radius=16, fill=CARD, outline=BORDER, width=2)
    draw.text((24, 18), label, font=FONT_PANEL, fill=TEXT)
    inner_x = (width - inner.width) // 2
    card.paste(inner, (inner_x, padding + label_h))
    return card


def render_vertical_collage(output: Path, title: str, panels: list[tuple[str, Path]], content_width: int = 1500) -> None:
    header_h = 120
    padding = 28
    panel_gap = 24
    rendered_panels: list[Image.Image] = []

    for label, path in panels:
        with Image.open(path) as src:
            image = src.convert("RGB")
            scale = min(1.0, content_width / image.width)
            resized = image.resize((int(image.width * scale), int(image.height * scale)), Image.LANCZOS)
        panel = add_panel_background(resized, label, content_width + 40)
        rendered_panels.append(panel)

    width = content_width + 96
    height = header_h + padding + sum(panel.height for panel in rendered_panels) + panel_gap * (len(rendered_panels) - 1) + padding
    canvas = Image.new("RGB", (width, height), BG)
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, 0, width, 18), fill=ACCENT)
    draw.text((34, 34), title, font=FONT_TITLE, fill=TEXT)

    y = header_h
    for panel in rendered_panels:
        x = (width - panel.width) // 2
        canvas.paste(panel, (x, y))
        y += panel.height + panel_gap

    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output)


def ensure_db_figures() -> dict[str, Path]:
    config = read_db_config()
    port_open = tcp_reachable(config["host"], int(config["port"]))
    table_count = run_scalar(config, "SELECT COUNT(*) AS table_count FROM information_schema.tables WHERE table_schema = 'public'")
    user_count = run_scalar(config, "SELECT COUNT(*) AS user_count FROM users")
    book_count = run_scalar(config, "SELECT COUNT(*) AS book_count FROM books")
    borrow_count = run_scalar(config, "SELECT COUNT(*) AS borrow_count FROM borrow_records")
    category_count = run_scalar(config, "SELECT COUNT(*) AS category_count FROM book_categories")
    log_count = run_scalar(config, "SELECT COUNT(*) AS log_count FROM operation_logs")
    backup_count = run_scalar(config, "SELECT COUNT(*) AS backup_count FROM backup_records")

    fig4_1 = FIG_DIR / "fig4_1_db_check.png"
    render_table_image(
        fig4_1,
        "数据库连接与库状态检测结果",
        "基于当前 db_config.txt 与 PostgreSQL 实例的实时检测结果",
        ["检测项", "结果"],
        [
            ["数据库主机", config["host"]],
            ["数据库端口", config["port"]],
            ["端口连通性", "成功" if port_open else "失败"],
            ["目标数据库", config["dbname"]],
            ["登录验证", "成功"],
            ["公共表数量", table_count],
        ],
        [320, 360],
        ["left", "left"],
    )

    fig4_2 = FIG_DIR / "fig4_2_summary.png"
    render_table_image(
        fig4_2,
        "核心数据汇总查询结果",
        "测试前系统核心业务表中的当前记录数量",
        ["统计指标", "数量"],
        [
            ["用户总数", user_count],
            ["图书总数", book_count],
            ["借阅记录数", borrow_count],
            ["分类总数", category_count],
            ["操作日志数", log_count],
            ["备份记录数", backup_count],
        ],
        [320, 180],
        ["left", "center"],
    )

    category_rows = run_copy_query(
        config,
        "SELECT category_id, category_name, COALESCE(parent_id::text, '-') AS parent_id FROM book_categories ORDER BY category_id",
    )
    fig4_8 = FIG_DIR / "fig4_8_categories.png"
    render_table_image(
        fig4_8,
        "图书分类数据查询结果",
        "book_categories 表中的分类结构数据",
        ["分类ID", "分类名称", "父分类ID"],
        [[row["category_id"], row["category_name"], row["parent_id"]] for row in category_rows],
        [140, 260, 180],
        ["center", "left", "center"],
    )

    log_rows = run_copy_query(
        config,
        "SELECT log_id, operator_id, operation_type, target_table, target_id, to_char(operation_time, 'YYYY-MM-DD HH24:MI:SS') AS operation_time FROM operation_logs ORDER BY log_id DESC LIMIT 10",
    )
    fig4_13 = FIG_DIR / "fig4_13_logs.png"
    render_table_image(
        fig4_13,
        "操作日志查询结果",
        "operation_logs 表最近的业务操作记录",
        ["日志ID", "操作人", "操作类型", "目标表", "目标ID", "操作时间"],
        [
            [
                row["log_id"],
                row["operator_id"],
                row["operation_type"],
                row["target_table"],
                row["target_id"],
                row["operation_time"],
            ]
            for row in log_rows
        ],
        [100, 100, 180, 180, 100, 260],
        ["center", "center", "left", "left", "center", "center"],
    )

    backup_rows = run_copy_query(
        config,
        "SELECT backup_id, backup_name, backup_type, operator_id, to_char(backup_time, 'YYYY-MM-DD HH24:MI:SS') AS backup_time FROM backup_records ORDER BY backup_id DESC LIMIT 10",
    )
    fig4_14 = FIG_DIR / "fig4_14_backups.png"
    render_table_image(
        fig4_14,
        "备份记录查询结果",
        "backup_records 表中的最近备份历史",
        ["备份ID", "备份文件名", "备份类型", "操作人", "备份时间"],
        [
            [row["backup_id"], row["backup_name"], row["backup_type"], row["operator_id"], row["backup_time"]]
            for row in backup_rows
        ],
        [100, 320, 160, 100, 260],
        ["center", "left", "left", "center", "center"],
    )

    setting_rows = run_copy_query(
        config,
        "SELECT setting_id, user_id, font_size, theme, language_pref, to_char(update_time, 'YYYY-MM-DD HH24:MI:SS') AS update_time FROM user_settings ORDER BY setting_id",
    )
    fig4_15 = FIG_DIR / "fig4_15_settings.png"
    render_table_image(
        fig4_15,
        "用户设置数据查询结果",
        "user_settings 表中的主题、字号与语言偏好",
        ["设置ID", "用户ID", "字号", "主题", "语言", "更新时间"],
        [
            [
                row["setting_id"],
                row["user_id"],
                row["font_size"],
                row["theme"],
                row["language_pref"],
                row["update_time"],
            ]
            for row in setting_rows
        ],
        [100, 100, 120, 120, 120, 260],
        ["center", "center", "center", "center", "center", "center"],
    )

    book_rows = run_copy_query(
        config,
        "SELECT book_id, book_title, available_stock, book_status FROM books ORDER BY book_id",
    )
    fig4_17_books = FIG_DIR / "fig4_17_books_after.png"
    render_table_image(
        fig4_17_books,
        "借阅后图书库存联动结果",
        "books 表中库存数量与图书状态的实时联动数据",
        ["图书ID", "书名", "可借库存", "图书状态"],
        [[row["book_id"], row["book_title"], row["available_stock"], row["book_status"]] for row in book_rows],
        [100, 300, 140, 180],
        ["center", "left", "center", "center"],
    )

    fig4_17 = FIG_DIR / "fig4_17_integration.png"
    render_vertical_collage(
        fig4_17,
        "图书借阅业务集成测试结果",
        [
            ("界面侧：借书成功提示", QT_IMG_DIR / "image4.png"),
            ("界面侧：借阅记录刷新", QT_IMG_DIR / "image6.png"),
            ("数据库侧：库存状态联动", fig4_17_books),
        ],
    )

    fig4_18 = FIG_DIR / "fig4_18_loglink.png"
    render_table_image(
        fig4_18,
        "日志联动测试结果",
        "新增图书与借阅办理业务触发的日志联动记录",
        ["操作类型", "目标表", "目标ID", "操作时间"],
        [[row["operation_type"], row["target_table"], row["target_id"], row["operation_time"]] for row in log_rows],
        [220, 220, 120, 280],
        ["left", "left", "center", "center"],
    )

    return {
        "fig4_1": fig4_1,
        "fig4_2": fig4_2,
        "fig4_8": fig4_8,
        "fig4_13": fig4_13,
        "fig4_14": fig4_14,
        "fig4_15": fig4_15,
        "fig4_16": FIG_DIR / "fig4_16_dashboard.png",
        "fig4_17": fig4_17,
        "fig4_18": fig4_18,
    }


def iter_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def find_table_after_caption(document: Document, caption_text: str) -> Table:
    blocks = list(iter_blocks(document))
    for index, block in enumerate(blocks):
        if isinstance(block, Paragraph) and block.text.strip() == caption_text:
            for next_block in blocks[index + 1:]:
                if isinstance(next_block, Table):
                    return next_block
            break
    raise ValueError(f"Unable to find table after caption: {caption_text}")


def set_run_font(run, size_pt: int, bold: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_paragraph_after(element, parent) -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    return Paragraph(new_p, parent)


def add_figure_block(anchor_element, parent, figure: FigureSpec):
    paragraph = add_paragraph_after(anchor_element, parent)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(figure.image_path), width=Cm(figure.width_cm))

    caption = add_paragraph_after(paragraph._element, parent)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run(figure.caption)
    set_run_font(caption_run, 11, bold=False)

    note = add_paragraph_after(caption._element, parent)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.first_line_indent = Cm(0.74)
    note_run = note.add_run(figure.note)
    set_run_font(note_run, 10, bold=False)

    spacer = add_paragraph_after(note._element, parent)
    return spacer._element


def build_figure_map(db_figures: dict[str, Path]) -> dict[str, list[FigureSpec]]:
    return {
        "表 4.5 系统通用类测试报告": [
            FigureSpec(
                db_figures["fig4_1"],
                "图4.1 数据库连接与库状态检测结果",
                "通过当前数据库配置完成端口、目标库和登录状态检测，结果表明 PostgreSQL 实例可正常连接，数据库基础环境满足系统测试条件。",
            ),
            FigureSpec(
                db_figures["fig4_2"],
                "图4.2 核心数据汇总查询结果",
                "对用户、图书、借阅、分类、日志和备份等核心业务表进行统计汇总，验证测试数据库已具备完整的基础测试数据。",
            ),
        ],
        "表 4.8 用户管理模块测试报告": [
            FigureSpec(
                QT_IMG_DIR / "image1.png",
                "图4.3 查询所有用户测试结果",
                "管理员进入用户管理界面后可正常加载全部用户列表，说明用户数据查询与分页展示功能运行正常。",
            ),
            FigureSpec(
                QT_IMG_DIR / "image2.png",
                "图4.4 按用户名查询用户测试结果",
                "输入指定用户名后系统能够返回对应用户记录，说明条件检索与结果过滤逻辑正确。",
            ),
            FigureSpec(
                QT_IMG_DIR / "image3.png",
                "图4.5 用户管理操作测试结果",
                "界面中各项用户管理操作按钮可正常触发，对应的编辑、删除等业务入口展示完整，满足管理员维护用户信息的需求。",
            ),
        ],
        "表 4.11 图书管理模块测试报告": [
            FigureSpec(
                QT_IMG_DIR / "image12.png",
                "图4.6 图书列表展示测试结果",
                "管理员进入图书管理界面后可正常查看图书基础信息、库存和状态，说明图书列表加载功能正常。",
            ),
            FigureSpec(
                QT_IMG_DIR / "image13.png",
                "图4.7 图书条件搜索测试结果",
                "输入关键字后系统能够按条件筛选图书记录，说明图书检索与结果刷新逻辑正确。",
            ),
            FigureSpec(
                db_figures["fig4_8"],
                "图4.8 图书分类数据查询测试结果",
                "在数据库侧查询图书分类数据，验证分类层级信息完整，可为前端分类选择和图书归类功能提供支撑。",
            ),
        ],
        "表 4.14 借阅业务模块测试报告": [
            FigureSpec(
                QT_IMG_DIR / "image4.png",
                "图4.9 借书成功测试结果",
                "读者完成借书操作后系统弹出成功提示，说明借阅业务的前端交互与提交流程正常。",
            ),
            FigureSpec(
                QT_IMG_DIR / "image5.png",
                "图4.10 还书确认测试结果",
                "执行还书操作时系统给出确认提示框，说明关键业务环节具备必要的二次确认控制。",
            ),
            FigureSpec(
                QT_IMG_DIR / "image6.png",
                "图4.11 还书完成后借阅记录刷新结果",
                "还书完成后借阅列表能够立即刷新并反映最新状态，说明借阅记录查询与界面联动正常。",
            ),
            FigureSpec(
                QT_IMG_DIR / "image7.png",
                "图4.12 借阅记录查询测试结果",
                "系统能够按照当前读者加载对应借阅记录，验证借阅历史查询功能可正常使用。",
            ),
        ],
        "表 4.17 系统维护模块测试报告": [
            FigureSpec(
                db_figures["fig4_13"],
                "图4.13 操作日志查询测试结果",
                "在数据库侧查询最近操作日志，可看到图书新增与借阅办理记录，说明关键业务动作已被正常追踪。",
            ),
            FigureSpec(
                db_figures["fig4_14"],
                "图4.14 备份记录查询测试结果",
                "查询备份记录表后能够获取历史备份文件与执行时间，说明系统维护模块的数据备份信息已正确落库。",
            ),
            FigureSpec(
                db_figures["fig4_15"],
                "图4.15 用户设置数据查询测试结果",
                "查询用户设置表可看到主题、字号和语言偏好等信息，说明系统配置项的持久化保存功能有效。",
            ),
        ],
        "表 4.18 系统集成测试用例": [
            FigureSpec(
                db_figures["fig4_16"],
                "图4.16 系统登录与主界面测试结果",
                "系统登录成功后可正常进入管理员主界面，并展示用户数、馆藏图书、未归还借阅和预约记录等汇总信息。",
                width_cm=15.2,
            ),
            FigureSpec(
                db_figures["fig4_17"],
                "图4.17 图书借阅业务集成测试结果",
                "借书成功后，界面提示、借阅记录刷新以及数据库库存状态变更保持一致，说明借阅业务在界面层与数据层之间联动正常。",
                width_cm=15.0,
            ),
            FigureSpec(
                db_figures["fig4_18"],
                "图4.18 日志联动测试结果",
                "对操作日志进行查询后，可看到新增图书与借阅办理两类业务记录，验证系统已具备基础的审计追踪能力。",
            ),
        ],
    }


def insert_figures_into_doc(figure_map: dict[str, list[FigureSpec]]) -> None:
    shutil.copy2(SOURCE_DOC, OUTPUT_DOC)
    document = Document(OUTPUT_DOC)
    for caption_text, figures in figure_map.items():
        table = find_table_after_caption(document, caption_text)
        current_element = table._element
        for figure in figures:
            if not figure.image_path.exists():
                raise FileNotFoundError(f"Missing image file: {figure.image_path}")
            current_element = add_figure_block(current_element, table._parent, figure)
    document.save(OUTPUT_DOC)


def verify_output(expected_added_images: int) -> tuple[int, int]:
    document = Document(OUTPUT_DOC)
    inline_shape_count = len(document.inline_shapes)
    figure_caption_count = sum(1 for paragraph in document.paragraphs if paragraph.text.strip().startswith("图4."))
    if figure_caption_count < expected_added_images:
        raise RuntimeError(
            f"Inserted caption count is lower than expected: {figure_caption_count} < {expected_added_images}"
        )
    return inline_shape_count, figure_caption_count


def main() -> None:
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"Source doc not found: {SOURCE_DOC}")
    if not PSQL_PATH.exists():
        raise FileNotFoundError(f"psql not found: {PSQL_PATH}")
    if not FONT_PATH.exists():
        raise FileNotFoundError(f"Font not found: {FONT_PATH}")

    db_figures = ensure_db_figures()
    figure_map = build_figure_map(db_figures)
    insert_figures_into_doc(figure_map)
    expected = sum(len(items) for items in figure_map.values())
    inline_count, caption_count = verify_output(expected)
    print(f"OUTPUT_DOC={OUTPUT_DOC}")
    print(f"INLINE_SHAPES={inline_count}")
    print(f"FIGURE_CAPTIONS={caption_count}")


if __name__ == "__main__":
    main()
